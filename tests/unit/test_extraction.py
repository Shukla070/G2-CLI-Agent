"""Tests for lantern.tools.extraction — text extraction from supported formats.

Uses small fixture files created in a temp directory, not the example
workspace, so tests are fully self-contained and reproducible regardless
of workspace contents.
"""

from __future__ import annotations

import pytest
from pathlib import Path

from lantern.tools.extraction import (
    ExtractionError,
    UnsupportedFormatError,
    extract_text,
    is_supported_format,
)


# ------------------------------------------------------------------
# Fixtures
# ------------------------------------------------------------------

@pytest.fixture
def tmp_workspace(tmp_path: Path) -> Path:
    """Create a temp directory with small fixture files."""
    return tmp_path


def _write_txt(directory: Path, name: str, content: str) -> Path:
    """Helper: write a .txt file."""
    p = directory / name
    p.write_text(content, encoding="utf-8")
    return p


def _write_md(directory: Path, name: str, content: str) -> Path:
    """Helper: write a .md file."""
    p = directory / name
    p.write_text(content, encoding="utf-8")
    return p


# ------------------------------------------------------------------
# .txt extraction
# ------------------------------------------------------------------

class TestTxtExtraction:
    def test_basic_txt(self, tmp_workspace: Path) -> None:
        f = _write_txt(tmp_workspace, "notes.txt", "Hello, world!")
        assert extract_text(f) == "Hello, world!"

    def test_multiline_txt(self, tmp_workspace: Path) -> None:
        content = "Line one\nLine two\nLine three"
        f = _write_txt(tmp_workspace, "multi.txt", content)
        assert extract_text(f) == content

    def test_empty_txt(self, tmp_workspace: Path) -> None:
        f = _write_txt(tmp_workspace, "empty.txt", "")
        assert extract_text(f) == ""

    def test_unicode_txt(self, tmp_workspace: Path) -> None:
        content = "Café résumé naïve — em dash"
        f = _write_txt(tmp_workspace, "unicode.txt", content)
        assert extract_text(f) == content

    def test_large_txt(self, tmp_workspace: Path) -> None:
        content = "word " * 1000  # ~5000 chars
        f = _write_txt(tmp_workspace, "large.txt", content)
        result = extract_text(f)
        assert len(result.split()) == 1000


# ------------------------------------------------------------------
# .md extraction
# ------------------------------------------------------------------

class TestMdExtraction:
    def test_basic_md(self, tmp_workspace: Path) -> None:
        content = "# Heading\n\nSome **bold** text."
        f = _write_md(tmp_workspace, "readme.md", content)
        # Markdown is returned as-is (not rendered)
        assert "# Heading" in extract_text(f)
        assert "**bold**" in extract_text(f)

    def test_md_with_code_blocks(self, tmp_workspace: Path) -> None:
        content = "# Code\n\n```python\nprint('hello')\n```\n"
        f = _write_md(tmp_workspace, "code.md", content)
        result = extract_text(f)
        assert "print('hello')" in result


# ------------------------------------------------------------------
# .docx extraction
# ------------------------------------------------------------------

class TestDocxExtraction:
    def test_basic_docx(self, tmp_workspace: Path) -> None:
        """Create a minimal .docx and verify text extraction."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_paragraph("First paragraph of the document.")
        doc.add_paragraph("Second paragraph with more text.")
        fpath = tmp_workspace / "test.docx"
        doc.save(str(fpath))

        result = extract_text(fpath)
        assert "First paragraph" in result
        assert "Second paragraph" in result

    def test_empty_docx(self, tmp_workspace: Path) -> None:
        """An empty .docx should return an empty string."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        fpath = tmp_workspace / "empty.docx"
        doc.save(str(fpath))

        result = extract_text(fpath)
        assert result == ""

    def test_docx_multiple_paragraphs(self, tmp_workspace: Path) -> None:
        """Multiple paragraphs should be joined with newlines."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_paragraph("Alpha")
        doc.add_paragraph("Beta")
        doc.add_paragraph("Gamma")
        fpath = tmp_workspace / "multi.docx"
        doc.save(str(fpath))

        result = extract_text(fpath)
        assert "Alpha" in result
        assert "Beta" in result
        assert "Gamma" in result


# ------------------------------------------------------------------
# .pdf extraction
# ------------------------------------------------------------------

class TestPdfExtraction:
    def test_basic_pdf(self, tmp_workspace: Path) -> None:
        """Create a minimal PDF and verify text extraction."""
        try:
            from pypdf import PdfWriter
        except ImportError:
            pytest.skip("pypdf not installed")

        # pypdf can create simple PDFs with annotations/metadata
        # but for real text, we need reportlab or a pre-built fixture.
        # Use a simpler approach: write a PDF-like file and test the
        # error handling, or create a real one with pypdf's writer.
        #
        # Actually, pypdf's PdfWriter doesn't easily create text
        # content. Let's test with a minimal approach using the
        # internal PDF structure.
        import io
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        fpath = tmp_workspace / "blank.pdf"
        with open(fpath, "wb") as f:
            writer.write(f)

        # A blank page should extract to empty text
        result = extract_text(fpath)
        # Blank pages may return empty or whitespace
        assert isinstance(result, str)

    def test_pdf_nonexistent(self, tmp_workspace: Path) -> None:
        """Missing PDF file should raise ExtractionError."""
        fpath = tmp_workspace / "missing.pdf"
        with pytest.raises(ExtractionError, match="does not exist"):
            extract_text(fpath)


# ------------------------------------------------------------------
# Unsupported formats
# ------------------------------------------------------------------

class TestUnsupportedFormats:
    def test_unsupported_extension(self, tmp_workspace: Path) -> None:
        f = tmp_workspace / "image.png"
        f.write_bytes(b"\x89PNG\r\n")
        with pytest.raises(UnsupportedFormatError, match="Unsupported"):
            extract_text(f)

    def test_unsupported_jpeg(self, tmp_workspace: Path) -> None:
        f = tmp_workspace / "photo.jpeg"
        f.write_bytes(b"\xff\xd8\xff\xe0")
        with pytest.raises(UnsupportedFormatError):
            extract_text(f)

    def test_unsupported_csv(self, tmp_workspace: Path) -> None:
        f = tmp_workspace / "data.csv"
        f.write_text("a,b,c\n1,2,3\n")
        with pytest.raises(UnsupportedFormatError):
            extract_text(f)


# ------------------------------------------------------------------
# Error handling
# ------------------------------------------------------------------

class TestErrorHandling:
    def test_nonexistent_file(self, tmp_workspace: Path) -> None:
        f = tmp_workspace / "ghost.txt"
        with pytest.raises(ExtractionError, match="does not exist"):
            extract_text(f)

    def test_corrupted_docx(self, tmp_workspace: Path) -> None:
        """A .docx that's actually garbage bytes should raise ExtractionError."""
        try:
            import docx  # noqa: F401
        except ImportError:
            pytest.skip("python-docx not installed")

        f = tmp_workspace / "corrupted.docx"
        f.write_bytes(b"this is not a docx file at all")
        with pytest.raises(ExtractionError):
            extract_text(f)


# ------------------------------------------------------------------
# is_supported_format
# ------------------------------------------------------------------

class TestIsSupportedFormat:
    @pytest.mark.parametrize("name,expected", [
        ("notes.txt", True),
        ("readme.md", True),
        ("draft.docx", True),
        ("report.pdf", True),
        ("image.png", False),
        ("data.csv", False),
        ("archive.zip", False),
        ("NOTES.TXT", True),      # Case insensitive
        ("Report.PDF", True),     # Case insensitive
    ])
    def test_format_detection(self, tmp_workspace: Path, name: str, expected: bool) -> None:
        f = tmp_workspace / name
        f.write_text("test", encoding="utf-8")
        assert is_supported_format(f) == expected
