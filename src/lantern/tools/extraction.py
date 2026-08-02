"""Document text extraction for Lantern.

Extracts plain text from the four supported document formats:
.txt, .md, .docx, and .pdf.  Scanned documents, images, and OCR
are explicitly out of scope per the assignment.

This module is format-aware but security-unaware — it receives
already-resolved ``Path`` objects (the sandbox check has already
happened by the time anything calls here) and returns strings.
Errors are wrapped in ``ExtractionError`` so callers get a
consistent interface regardless of which library failed underneath.

Libraries:
    .txt / .md  — built-in ``Path.read_text()``
    .docx       — ``python-docx`` (``Document.paragraphs``)
    .pdf        — ``pypdf`` (``PdfReader.pages[].extract_text()``)
"""

from __future__ import annotations

from pathlib import Path


# ------------------------------------------------------------------
# Exceptions
# ------------------------------------------------------------------

class ExtractionError(Exception):
    """Any failure during text extraction.

    Wraps the underlying library error so callers don't need to know
    whether it was a ``docx`` parse error or a ``pypdf`` read error.
    """


class UnsupportedFormatError(ExtractionError):
    """Raised when the file extension has no extraction backend."""


# ------------------------------------------------------------------
# Format handlers (private)
# ------------------------------------------------------------------

_SUPPORTED_EXTENSIONS: frozenset[str] = frozenset({".txt", ".md", ".docx", ".pdf"})


def _extract_plaintext(path: Path) -> str:
    """Handle .txt and .md — read as UTF-8 text."""
    try:
        return path.read_text(encoding="utf-8", errors="replace")
    except OSError as exc:
        raise ExtractionError(f"Failed to read {path.name}: {exc}") from exc


def _extract_docx(path: Path) -> str:
    """Handle .docx — extract paragraph text via python-docx."""
    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError as exc:
        raise ExtractionError(
            "python-docx is required for .docx extraction but is not installed."
        ) from exc

    try:
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except Exception as exc:
        raise ExtractionError(
            f"Failed to extract text from {path.name}: {exc}"
        ) from exc


def _extract_pdf(path: Path) -> str:
    """Handle .pdf — extract page text via pypdf."""
    try:
        from pypdf import PdfReader  # type: ignore[import-untyped]
    except ImportError as exc:
        raise ExtractionError(
            "pypdf is required for .pdf extraction but is not installed."
        ) from exc

    try:
        reader = PdfReader(str(path))
        pages: list[str] = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n".join(pages)
    except Exception as exc:
        raise ExtractionError(
            f"Failed to extract text from {path.name}: {exc}"
        ) from exc


# Extension → handler mapping
_HANDLERS: dict[str, callable] = {
    ".txt": _extract_plaintext,
    ".md": _extract_plaintext,
    ".docx": _extract_docx,
    ".pdf": _extract_pdf,
}


# ------------------------------------------------------------------
# Public API
# ------------------------------------------------------------------

def extract_text(path: Path) -> str:
    """Extract plain text from a supported document file.

    Parameters
    ----------
    path : Path
        Fully resolved path to the file.  Must already have passed
        through ``Sandbox.resolve()`` — this function does no
        sandbox validation itself.

    Returns
    -------
    str
        The extracted text content.  May be empty if the document
        contains no extractable text (e.g. a blank page PDF).

    Raises
    ------
    UnsupportedFormatError
        If the file extension is not one of .txt, .md, .docx, .pdf.
    ExtractionError
        If the file cannot be read or parsed.
    """
    if not path.exists():
        raise ExtractionError(f"File does not exist: {path.name}")

    ext = path.suffix.lower()
    handler = _HANDLERS.get(ext)
    if handler is None:
        raise UnsupportedFormatError(
            f"Unsupported file format '{ext}'. "
            f"Supported formats: {', '.join(sorted(_SUPPORTED_EXTENSIONS))}"
        )

    return handler(path)


def is_supported_format(path: Path) -> bool:
    """Return True if the file extension has an extraction backend."""
    return path.suffix.lower() in _SUPPORTED_EXTENSIONS
